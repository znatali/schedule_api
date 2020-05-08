from django.http import HttpResponse
from docx import Document
from rest_framework import exceptions, mixins, permissions, viewsets
from rest_framework.decorators import action
from rest_framework.permissions import AllowAny

from server.apps.main.logic.schedule_day_serializer import ScheduleDaySerializer
from server.apps.main.logic.schedule_item_serializer import (
    ScheduleItemCreateSerializer,
    ScheduleItemSerializer,
    ScheduleItemUpdateSerializer,
)
from server.apps.main.logic.schedule_serializer import ScheduleSerializer
from server.apps.main.models.schedule import Schedule
from server.apps.main.models.schedule_day import ScheduleDay
from server.apps.main.models.schedule_item import ScheduleItem


class ScheduleViewSet(
    mixins.CreateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.UpdateModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    viewsets.GenericViewSet,
):
    """Updates and retrieves schedules."""

    queryset = Schedule.objects.all()
    serializer_class = ScheduleSerializer
    permission_classes = (permissions.IsAuthenticated,)

    @action(detail=True, methods=['get'])
    def download(self, request, pk):
        """Get formatted schedule."""
        document = Document()
        schedule = self.get_object()
        document.add_heading(schedule.title, 0)
        document.add_paragraph(schedule.faculty.title)
        document.add_paragraph(f'Курс {str(schedule.year)}')
        document.add_paragraph(f'Семестр {str(schedule.term)}')
        document.add_paragraph(f'Форма обучения {schedule.education_format}')

        schedule_days = ScheduleDay.objects.filter(schedule=schedule).all()

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'День'
        hdr_cells[1].text = 'Время'
        hdr_cells[2].text = 'Занятие'

        for schedule_day in schedule_days:
            row_cells = table.add_row().cells
            row_cells[0].text = f'{schedule_day.title} ' +\
                f'{schedule_day.date.strftime("%d-%b-%Y")}'  # noqa: WPS323
            schedule_items = ScheduleItem.objects.filter(schedule_day=schedule_day).all()
            for schedule_item in schedule_items:
                row_cells = table.add_row().cells
                row_cells[1].text = \
                    f'{schedule_item.start_time.strftime("%H:%M")} - ' \
                    + f'{schedule_item.end_time.strftime("%H:%M")}'
                row_cells[2].text = \
                    f'{schedule_item.title} ({schedule_item.type}) {schedule_item.teacher.get_full_name()}'

        http_word_response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        document.save(http_word_response)
        http_word_response['Content-Disposition'] = 'attachment; filename=schedule.docx'
        return http_word_response


class ScheduleDayViewSet(
    mixins.CreateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.UpdateModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    viewsets.GenericViewSet,
):
    """Updates and retrieves schedules for a day."""

    queryset = ScheduleDay.objects.all()
    serializer_class = ScheduleDaySerializer
    permission_classes = (permissions.IsAuthenticated,)


class ScheduleItemViewSet(
    mixins.CreateModelMixin,
    mixins.RetrieveModelMixin,
    mixins.UpdateModelMixin,
    mixins.ListModelMixin,
    mixins.DestroyModelMixin,
    viewsets.GenericViewSet,
):
    """Updates and retrieves schedule items."""

    queryset = ScheduleItem.objects.all()
    serializer_class = ScheduleItemSerializer
    permission_classes = (permissions.IsAuthenticated,)

    def get_serializer(self, *args, **kwargs):
        """
        Route serializers based on action.

            list                    ScheduleItemSerializer
            retrieve                ScheduleItemSerializer
            update, partial_update  ScheduleItemUpdateSerializer
            create                  ScheduleItemCreateSerializer
        """
        if self.request.user.is_anonymous and AllowAny in self.permission_classes:
            kwargs['context'] = self.get_serializer_context()
            return self.serializer_class(*args, **kwargs)
        if self.action in {'list', 'retrieve'}:
            serializer_class = ScheduleItemSerializer
        elif self.action in {'update', 'partial_update'}:
            serializer_class = ScheduleItemUpdateSerializer
        elif self.action == 'create':
            serializer_class = ScheduleItemCreateSerializer
        else:
            serializer_class = None

        if serializer_class is None:
            raise exceptions.ValidationError(detail={self.action: ['Permission error.']})
        kwargs['context'] = self.get_serializer_context()
        return serializer_class(*args, **kwargs)
